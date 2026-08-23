"""
Minimal python-docx helpers for building the paper.

Kept separate from build_paper.py so the paper module is prose and numbers, not
formatting boilerplate.
"""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"


def setup_styles(doc) -> None:
    """Two-column-ready single-column layout with conventional paper typography."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, bold in [("Heading 1", 13, True), ("Heading 2", 11.5, True),
                             ("Heading 3", 10.5, True)]:
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(4)


def title_block(doc, title: str, authors: str, affiliation: str, note: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    for text, size, italic in [(authors, 11, False), (affiliation, 10, True),
                               (note, 9, True)]:
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.italic = italic


def abstract(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Abstract—")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(10)


def para(doc, text: str, style: str | None = None):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def code_block(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = MONO_FONT
    r.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)


def table(doc, caption: str, headers: list[str], rows: list[list[str]],
          bold_rows: set[int] | None = None) -> None:
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(2)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)

    bold_rows = bold_rows or set()
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(8.5)
            if ri in bold_rows:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(8)
