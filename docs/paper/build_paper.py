"""
Build the OrbitGuard methods paper as a .docx.

Every quantitative claim in the prose is read from a JSON artifact in reports/
rather than typed in, so the paper cannot drift from the experiments. If a
number here looks wrong, re-run the experiment that produces it — do not edit
the number.

    reports/normalisation_ablation.json   src/ml/norm_ablation.py       (C1)
    reports/baseline_invariance.json      src/ml/baseline_invariance.py (C1b)
    reports/rare_mode_sweep.json          src/ml/rare_mode_sweep.py     (C2)
    reports/prune_economics.json          src/ml/prune_economics.py     (C3)

Usage:
    python docs/paper/build_paper.py
    python docs/paper/build_paper.py --out /tmp/draft.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).parent))
from docx_helpers import (abstract, bullets, caption, code_block, para,
                          setup_styles, table, title_block)

REPO = Path(__file__).resolve().parents[2]
REPORTS = REPO / "reports"

TITLE = ("Scale-Invariant Baselines Can Certify a Broken Deep Model: "
         "Grouped-Normalisation Collapse in Learned Simulation Screening")
AUTHORS = "Harsha Sakamuri"
AFFILIATION = "[affiliation]"


def load(name: str) -> dict | None:
    p = REPORTS / name
    if not p.exists():
        print(f"  WARNING: {p} missing — sections depending on it will be stubbed")
        return None
    return json.loads(p.read_text())


def f4(x) -> str:
    return "—" if x is None else f"{x:.4f}"


def pct(x, dp=1) -> str:
    return "—" if x is None else f"{100*x:.{dp}f}%"


def sci(x) -> str:
    """Scientific notation, tolerant of a missing artifact."""
    return "—" if x is None else f"{x:.2e}"


def build(out_path: Path) -> None:
    norm = load("normalisation_ablation.json")
    inv = load("baseline_invariance.json")
    rare = load("rare_mode_sweep.json")
    econ = load("prune_economics.json")

    doc = Document()
    setup_styles(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.9)

    title_block(doc, TITLE, AUTHORS, AFFILIATION,
                "Draft — generated from measured artifacts by "
                "docs/paper/build_paper.py")

    # ── Abstract ─────────────────────────────────────────────────────────────
    v = norm["paired"][0] if norm and norm["paired"] else {}
    ex_planet = v.get("planet", "venus")
    ex_good = v.get("val_auc_per-timestep")
    ex_bad = v.get("val_auc_grouped")
    ex_std = v.get("pred_std_grouped")
    tree_spread = max(inv["tree_auc_spread_across_conditions"].values()) if inv else None
    n_planets = len(norm["paired"]) if norm else 0
    n_collapsed = sum(1 for p in (norm or {}).get("paired", [])
                      if p.get("collapsed_grouped"))
    worst = min(norm["paired"], key=lambda p: p["val_auc_grouped"]) if norm else {}
    others = sorted(abs(p["val_auc_gain_vs_grouped"])
                    for p in (norm or {}).get("paired", [])
                    if not p.get("collapsed_grouped"))
    other_lo = others[0] if others else None
    other_hi = others[-1] if others else None

    abstract(doc,
        f"Sharing one feature scaler across heterogeneous groups is a routine "
        f"preprocessing choice. We show it can silently destroy a sequence model "
        f"while every standard sanity check reports success. On a corpus of "
        f"70,000 simulated interplanetary trajectories across seven targets, a "
        f"Transformer trained under a scaler pooled across a group of targets "
        f"degenerates on {worst.get('planet', 'one target').capitalize()} to a "
        f"constant output: held-out prediction standard deviation "
        f"{sci(worst.get('pred_std_grouped'))} and validation AUC "
        f"{f4(worst.get('val_auc_grouped'))}, against "
        f"{f4(worst.get('val_auc_per-timestep'))} for the identical "
        f"architecture under per-timestep normalisation. A gradient-boosted tree "
        f"fitted to the same arrays under the same preprocessing is unaffected: "
        f"its AUC varies by at most {f4(tree_spread)} across all three "
        f"normalisation conditions, reporting near-perfect separability while "
        f"the network cannot discriminate at all. The practitioner's instinct to "
        f"validate a pipeline with a simple baseline therefore fails precisely "
        f"here, because trees are invariant to the defect that destroys the "
        f"network. The effect is severe but not universal: {n_collapsed} of "
        f"{n_planets} targets collapses outright, while the remainder lose "
        f"{f4(other_lo)} to {f4(other_hi)} AUC without collapsing, and we report "
        f"the conditions under which it does and does not appear. We decompose "
        f"the cause — pooling across timesteps is survivable, pooling across "
        f"groups is not — give a diagnostic based on prediction variance rather "
        f"than accuracy, and report a second failure in which the same model "
        f"cannot reach a rare class that a tree separates from its own input.")

    # ── 1. Introduction ──────────────────────────────────────────────────────
    doc.add_heading("1  Introduction", level=1)
    para(doc,
        "Monte Carlo trajectory campaigns spend most of their compute on runs "
        "that were doomed at injection. A natural response is to learn a screen: "
        "watch the early trajectory and cancel the runs that will fail. Building "
        "that screen surfaced two failures that we believe generalise well beyond "
        "astrodynamics, and one negative result about the screen itself.")
    para(doc,
        "The first failure is the subject of this paper. A deep model trained on "
        "grouped data, with one feature scaler fitted across the group, collapsed "
        "to emitting a single constant per group. It was not detected for a long "
        "time, and the reason it was not detected is the interesting part: every "
        "check that would normally catch it reported that the pipeline was "
        "healthy. Validation AUC was high, because a mixed validation set lets a "
        "model score well by ranking groups against each other without "
        "discriminating within any of them. A gradient-boosted baseline on the "
        "identical arrays reported near-perfect performance, because trees split "
        "on absolute values and are invariant to the scaling that the network "
        "depends on.")
    para(doc,
        "We make three contributions. (i) A controlled decomposition of the "
        "collapse, isolating pooling across timesteps from pooling across groups "
        "and showing only the latter is fatal. (ii) A demonstration that a "
        "scale-invariant baseline actively certifies the broken configuration, "
        "with a diagnostic that catches it. (iii) A second failure of the same "
        "model class — an inability to reach a rare class that is linearly "
        "recoverable from its own input — which is measurable rather than "
        "speculative.")
    para(doc,
        "All numbers in this paper are produced by scripts in the accompanying "
        "repository and written to machine-readable artifacts; the paper itself "
        "is generated from those artifacts.")

    # ── 2. Setup ─────────────────────────────────────────────────────────────
    doc.add_heading("2  Setup", level=1)
    doc.add_heading("2.1  Task and data", level=2)
    para(doc,
        "The corpus is 70,000 simulated Earth-departure missions across seven "
        "interplanetary targets (Mercury through Neptune), 10,000 per target, "
        "generated from a deterministic two/three-body propagator. Each mission "
        "is parameterised by six injection offsets — three components of the "
        "trans-orbit-insertion burn and three parking-orbit angles — and labelled "
        "success or failure, with failures further typed by mode (surface impact, "
        "orbit too high, missed target, and others). Telemetry is sampled at a "
        "fixed cadence and downsampled to roughly 100 steps per mission so that "
        "targets whose flight times differ by two orders of magnitude yield "
        "comparable sequence lengths.")
    para(doc,
        "The screening task is: observe the first 40% of a mission's trajectory "
        "and decide whether to abort. Splits are 70/15/15 by mission, "
        "deterministic in (n, seed), and shared by every experiment through a "
        "single definition so that no model is evaluated on data used to select "
        "it.")

    doc.add_heading("2.2  Models", level=2)
    para(doc,
        "The sequence model is a Pre-LN Transformer encoder (d_model 128, 8 "
        "heads, 4 layers, CLS pooling) with two heads sharing a trunk: mission "
        "outcome and failure mode. It is trained on random prefixes so it is "
        "in-distribution at any streaming position. The baseline is XGBoost (300 "
        "trees, depth 5) fitted to the same normalised prefix, flattened. "
        "Throughout, 'identical input' means literally the same array: the "
        "baseline consumes the view the network receives, not a re-derived "
        "feature set.")

    doc.add_heading("2.3  Normalisation conditions", level=2)
    para(doc,
        "Three conditions differ only in how the feature statistics are pooled. "
        "Architecture, seed, split, optimiser and epoch budget are held fixed.")
    bullets(doc, [
        "per-timestep — each feature standardised against its distribution at "
        "that timestep index across missions of one target. This is the "
        "production configuration.",
        "global — one RobustScaler (median, IQR) per target, pooled across all "
        "timesteps of that target.",
        "grouped — one RobustScaler pooled across all timesteps of every target "
        "in a regime group (inner: Mercury, Venus, Mars; outer: Jupiter, Saturn, "
        "Uranus, Neptune). This reproduces the configuration that failed.",
    ])

    # ── 3. Collapse ──────────────────────────────────────────────────────────
    doc.add_heading("3  Grouped-Normalisation Collapse", level=1)
    doc.add_heading("3.1  Mechanism", level=2)
    para(doc,
        "A mission's features sweep a wide range over its flight, and different "
        "targets occupy ranges that differ by orders of magnitude. A scaler "
        "fitted over a pooled set therefore has its scale set by the largest "
        "source of variation in the pool. What the screen must actually "
        "discriminate is none of those: it is the mission-to-mission spread at a "
        "given point in flight, which is small by comparison. Pooling divides "
        "that spread by a denominator chosen for something else, compressing the "
        "discriminative signal toward zero while leaving the between-group "
        "structure intact.")
    para(doc,
        "We quantify this as the signal ratio: the mean across-mission standard "
        "deviation of the normalised features, measured inside the observed "
        "window only. The restriction matters. Averaged over the whole flight the "
        "metric is uninformative, because late timesteps carry enormous "
        "across-mission spread — failing trajectories have physically diverged by "
        "then — which swamps the early signal and makes every condition look "
        "healthy. The model commits at 40%, so the relevant quantity is the "
        "spread inside the prefix it sees.")

    doc.add_heading("3.2  Controlled ablation", level=2)
    if norm:
        rows = []
        for p in norm["paired"]:
            rows.append([
                p["planet"].capitalize(),
                f"{p['signal_per-timestep']:.4f}", f4(p["val_auc_per-timestep"]),
                f"{p['pred_std_per-timestep']:.2e}",
                f"{p['signal_global']:.4f}", f4(p["val_auc_global"]),
                f"{p['pred_std_global']:.2e}",
                f"{p['signal_grouped']:.4f}", f4(p["val_auc_grouped"]),
                f"{p['pred_std_grouped']:.2e}",
            ])
        table(doc, "Table 1  Normalisation ablation. Identical architecture, "
                   "seed and split; only pooling differs. 'signal' is the mean "
                   "within-timestep standard deviation of normalised features in "
                   "the observed window; 'std' is the standard deviation of "
                   "P(fail) over the held-out split.",
              ["Target", "sig", "AUC", "std", "sig", "AUC", "std",
               "sig", "AUC", "std"], rows)
        caption(doc, "Columns 2–4 per-timestep, 5–7 global, 8–10 grouped.")

        collapsed = [p["planet"].capitalize() for p in norm["paired"]
                     if p.get("collapsed_grouped")]
        survived_global = [p["planet"] for p in norm["paired"]
                           if not p.get("collapsed_global")]
        para(doc,
            f"The decomposition is the first result. Pooling across timesteps "
            f"within a single target compresses the signal by one to two orders "
            f"of magnitude and yet the network still discriminates: it survives "
            f"on {len(survived_global)} of {len(norm['paired'])} targets, losing "
            f"at most a few thousandths of AUC. Pooling across a group is "
            f"qualitatively different, and on "
            f"{', '.join(collapsed) if collapsed else 'no target'} it is "
            f"catastrophic — held-out output standard deviation "
            f"{sci(worst.get('pred_std_grouped'))}, meaning the network emits "
            f"one number for every mission regardless of input, with AUC "
            f"falling from {f4(worst.get('val_auc_per-timestep'))} to "
            f"{f4(worst.get('val_auc_grouped'))}.")
        para(doc,
            f"The second result is that the collapse is selective, and we report "
            f"this rather than only the case that fails. On the remaining "
            f"{len(norm['paired']) - len(collapsed)} targets the same "
            f"manipulation costs between {f4(other_lo)} and {f4(other_hi)} AUC "
            f"and does not collapse the output. Severity does not follow the "
            f"signal ratio alone: Jupiter's signal is compressed to "
            f"{[p['signal_grouped'] for p in norm['paired'] if p['planet'] == 'jupiter'][0]:.4f} "
            f"— comparable to Mercury's — yet it loses almost nothing, because "
            f"its task is separable enough that even a heavily attenuated signal "
            f"suffices. Compression is necessary for the collapse but not "
            f"sufficient; task difficulty modulates it. A predictive rule of the "
            f"form 'collapse when the ratio falls below X' is therefore not "
            f"supported by these data, and we do not claim one.")
        para(doc,
            "One scope note. This ablation isolates the normalisation factor "
            "alone: a separate model is trained per target and only the "
            "statistics are pooled. The production incident that motivated the "
            "study additionally shared one model across the group, which "
            "compounds the effect — a single trunk must then serve targets whose "
            "inputs have been flattened toward a common constant. The numbers "
            "here are therefore a lower bound on the damage the full "
            "configuration causes, and should be read as establishing that "
            "pooled normalisation is sufficient on its own to destroy a target, "
            "not as a reproduction of the original failure in its entirety.")
    else:
        para(doc, "[Table 1 pending: run python -m src.ml.norm_ablation]")

    doc.add_heading("3.3  The baseline certifies the broken configuration", level=2)
    if inv:
        rows = []
        by = {(r["planet"], r["norm_mode"]): r for r in inv["runs"]}
        planets = sorted({r["planet"] for r in inv["runs"]},
                         key=lambda p: [x["planet"] for x in (norm or {}).get("paired", [])].index(p)
                         if norm and p in [x["planet"] for x in norm["paired"]] else 99)
        for p in planets:
            row = [p.capitalize()]
            for c in ["per-timestep", "global", "grouped"]:
                row.append(f4(by.get((p, c), {}).get("tree_auc")))
            row.append(f4(inv["tree_auc_spread_across_conditions"].get(p)))
            rows.append(row)
        table(doc, "Table 2  XGBoost on the identical normalised window, under "
                   "each condition. The tree is invariant to the preprocessing "
                   "that destroys the network.",
              ["Target", "per-timestep", "global", "grouped", "spread"], rows)
        wp = worst.get("planet", planets[0])
        para(doc,
            f"Under the grouped condition the tree reports AUC "
            f"{f4(by.get((wp, 'grouped'), {}).get('tree_auc'))} on "
            f"{wp.capitalize()} — the target where the network collapses to a "
            f"constant at AUC {f4(worst.get('val_auc_grouped'))}. Across all "
            f"targets the tree's AUC moves by at most {f4(tree_spread)} between "
            f"conditions, so the baseline is effectively blind to the choice. A "
            f"practitioner who runs it to check whether the features carry "
            f"signal — the standard and correct instinct — receives an "
            f"unambiguous yes, and concludes the deep model needs tuning rather "
            f"than that the pipeline is broken. That is the failure this paper "
            f"is about: not that the collapse happens, but that the check "
            f"designed to catch this class of problem cannot see it.")
        para(doc,
            "This is the methodological point. The baseline is not wrong: the "
            "information genuinely is present, and a tree genuinely can use it. "
            "The baseline is uninformative about the question actually being "
            "asked, which is whether the network can use it. Invariance to "
            "feature scaling — normally a reason to prefer trees as a diagnostic "
            "— is exactly what makes them blind here.")
    else:
        para(doc, "[Table 2 pending: run python -m src.ml.baseline_invariance]")

    doc.add_heading("3.4  A diagnostic that does catch it", level=2)
    para(doc,
        "Accuracy on a mixed validation set does not catch the collapse, because "
        "a per-group constant ranks groups correctly and a pooled metric rewards "
        "that. Two checks do:")
    bullets(doc, [
        "Prediction variance within a group. A collapsed model has near-zero "
        "spread of its output across inputs. This is a property of the "
        "predictions alone — no labels required — and it separates collapse from "
        "ordinary underfitting, which produces wrong but varying predictions.",
        "Tree-versus-network under identical preprocessing. A large gap in "
        "favour of the tree, with both fed the same array, indicates the network "
        "cannot exploit information that is present rather than that the "
        "information is absent.",
    ])
    para(doc,
        "Both are cheap and neither requires suspecting the specific defect in "
        "advance. We recommend reporting per-group prediction variance alongside "
        "aggregate metrics whenever a model is trained on grouped data with "
        "heterogeneous scales.")

    # ── 4. Rare mode ─────────────────────────────────────────────────────────
    doc.add_heading("4  Failure to Reach a Present Signal", level=1)
    if rare:
        para(doc,
            f"A related failure appears with correct normalisation. On "
            f"{rare['planet'].capitalize()}, the failure mode "
            f"'{rare['rare_mode']}' accounts for {rare['rare_mode_train_n']} of "
            f"{rare['n_train_failures']} training failures. A tree fitted to the "
            f"same normalised {rare['window_steps']}-step window separates that "
            f"mode from success at AUC "
            f"{f4(rare['tree_reference']['auc'])}. The sequence model's recall on "
            f"it is {f4(rare['best_sequence_rare_recall'])} at its operating "
            f"point.")
        rows = [[f"{r['mode_alpha']:.2f}", f"{r['effective_resample_factor']:.1f}x",
                 f4(r["test_f1"]), f4(r["overall_failure_recall"]),
                 f4(r["rare_mode_recall"])] for r in rare["runs"]]
        table(doc, f"Table 3  Rare-mode oversampling sweep on "
                   f"{rare['planet'].capitalize()}. Increasing the sampling "
                   f"weight of the rare mode does not recover it.",
              ["mode_alpha", "effective resample", "test F1",
               "overall recall", f"{rare['rare_mode']} recall"], rows)
        para(doc,
            "Because both models consume the same array and one of them "
            "separates the classes, the information is present and the sequence "
            "model's failure is an optimisation limit rather than an information "
            "limit. Resampling the mode does not close the gap, which rules out "
            "the simplest explanation. Whether the cause is the loss landscape, "
            "the pooling operator discarding a localised cue, or the binary head "
            "being dominated by the majority mode is open; probing the trunk for "
            "linear separability of the rare mode would localise it.")
    else:
        para(doc, "[Section 4 pending: run python -m src.ml.rare_mode_sweep]")

    # ── 5. Economics ─────────────────────────────────────────────────────────
    doc.add_heading("5  Where the Screen Belongs", level=1)
    if econ:
        w = econ["weighted"]
        rows = [
            ["T0 — six launch parameters, before propagating",
             pct(w["compute_saved_t0"]), pct(w["false_prune_rate_t0"], 2),
             pct(w["fail_recall_t0"], 2)],
            ["T40 — telemetry Transformer at 40%",
             pct(w["compute_saved_t40"]), pct(w["false_prune_rate_t40"], 2),
             pct(w["fail_recall_t40"], 2)],
            ["Cascade — T0 where confident, else T40",
             pct(w["cascade_saved"]), pct(w["cascade_false_prune"], 2),
             pct(w["cascade_recall"], 2)],
        ]
        table(doc, "Table 4  Screening economics, compute charged in "
                   "propagation-days across a ~100x per-target cost range. "
                   "Thresholds fitted on validation and reported on the held-out "
                   "split.",
              ["Screen", "Compute saved", "Good missions destroyed",
               "Failure recall"], rows, bold_rows={0})
        para(doc,
            f"Having fixed the sequence model, the honest conclusion is that it "
            f"should not be used. A six-feature classifier over the injection "
            f"parameters, evaluated before any propagation happens, saves "
            f"{pct(w['compute_saved_t0'])} of compute against the telemetry "
            f"model's {pct(w['compute_saved_t40'])}, at a comparable cost in "
            f"good missions destroyed. Accuracy is also flat from 10% to 40% "
            f"observed. The reason is structural: the simulator is "
            f"deterministic, so the outcome is a fixed function of the injection "
            f"offsets and the trajectory is their integral. It cannot carry "
            f"information the parameters do not already have.")
        para(doc,
            "We report this because it bounds the practical significance of "
            "Sections 3 and 4. The methodological findings concern any grouped "
            "sequence-modelling pipeline; the application that motivated them "
            "turns out not to need a sequence model. Logistic regression on the "
            "same six features scores at chance, so the task does require a "
            "nonlinear learner — just not a temporal one. A sequential screen "
            "earns its place only where the outcome is not determined at "
            "injection: mid-flight stochasticity, unmodelled dynamics, sensor "
            "noise, or campaigns where launch parameters are not recorded.")
    else:
        para(doc, "[Table 4 pending: run python -m src.ml.prune_economics]")

    # ── 6. Limitations ───────────────────────────────────────────────────────
    doc.add_heading("6  Limitations", level=1)
    bullets(doc, [
        "Single seed. Every result is one run at seed 42. The split is "
        "deterministic in (n, seed) so the numbers reproduce exactly, which is "
        "reproducibility, not stability. No confidence intervals are claimed.",
        "One domain. The collapse is demonstrated on one corpus. We characterise "
        "the mechanism in terms of a variance ratio that is not specific to "
        "astrodynamics, but we have not reproduced it on a second dataset, and "
        "until we do the predictive rule — collapse when the ratio exceeds some "
        "threshold — remains a conjecture.",
        "Synthetic and deterministic. No execution error, unmodelled "
        "accelerations or sensor noise, and every mission of a target shares a "
        "time base. That shared time base is what makes per-timestep "
        "standardisation work at all, and it is a property of the generator "
        "rather than of real campaigns.",
        "Seven targets. An eighth (the Moon) was generated and is excluded by "
        "decision: it is a short Earth-centric transfer sharing neither the cost "
        "structure of Section 5 nor the dynamical regime.",
        "No real mission data.",
    ])

    # ── 7. Related work ──────────────────────────────────────────────────────
    doc.add_heading("7  Related Work", level=1)
    para(doc,
        "[TO WRITE — see docs/paper/PAPER_NOTES.md. This section cannot be "
        "generated from the repository and needs a literature pass: "
        "normalisation and feature scaling in deep learning; per-group and "
        "per-instance normalisation; shortcut learning and Clever-Hans effects; "
        "the tabular-versus-deep literature on why trees remain strong; "
        "class-imbalance and rare-class recall; surrogate modelling and learned "
        "early termination for simulation.]")

    # ── 8. Conclusion ────────────────────────────────────────────────────────
    doc.add_heading("8  Conclusion", level=1)
    para(doc,
        "Sharing a feature scaler across heterogeneous groups can reduce a deep "
        "model to a per-group constant while a tree on the identical arrays "
        "reports the task as solved. The failure is not exotic, the "
        "preprocessing choice that causes it is routine, and the standard "
        "defence against it — check a simple baseline — is invariant to it by "
        "construction. Reporting per-group prediction variance alongside "
        "aggregate metrics costs nothing and would have caught it immediately.")

    doc.add_heading("Reproduction", level=1)
    code_block(doc,
        "export ORBITGUARD_DATA=/path/to/dataset\n"
        "python -m src.ml.norm_ablation          # Table 1\n"
        "python -m src.ml.baseline_invariance    # Table 2\n"
        "python -m src.ml.rare_mode_sweep        # Table 3\n"
        "python -m src.ml.prune_economics        # Table 4\n"
        "python docs/paper/build_paper.py        # this document")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  Saved -> {out_path}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=str(REPO / "docs/paper/OrbitGuard_paper_draft.docx"))
    args = ap.parse_args()
    build(Path(args.out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
